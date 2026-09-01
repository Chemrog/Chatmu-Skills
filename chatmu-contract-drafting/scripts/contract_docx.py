#!/usr/bin/env python3
"""Draft a music contract (.docx) from a JSON brief.

Input: /workspace/in/contract.json (or first argv) with:
  contract_type: "split_sheet" | "producer" | "management" | "feature"
  + parties / terms. Produces a draft with placeholders for lawyer review.

Convention: output /workspace/out/<type>_draft_<n>.docx.
Run only via the execute_python sandbox. python-docx preinstalled.
"""
import json
import os
import sys

from docx import Document
from docx.shared import Pt


def load_input() -> dict:
    path = sys.argv[1] if len(sys.argv) > 1 else "/workspace/in/contract.json"
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")


def split_sheet(doc, c):
    doc.add_heading("Split Sheet (Draft for Review)", level=1)
    for song in c.get("songs", [c]):
        doc.add_paragraph(f"Song: {song.get('title', 'TBD')}")
        doc.add_paragraph(f"ISRC (if released): {song.get('isrc', 'TBD')}")
        doc.add_paragraph("")
        p = doc.add_paragraph()
        r = p.add_run("Composition split (writers must total 50%, publishers 50%):")
        r.bold = True
        for w in song.get("writers", []):
            doc.add_paragraph(f"- {w.get('name', '')} — Writer {w.get('share_pct', 0)}% (PRO: {w.get('pro', '')} / IPI: {w.get('ipi', '')})")
        for pbl in song.get("publishers", []):
            doc.add_paragraph(f"- {pbl.get('name', '')} — Publisher {pbl.get('share_pct', 0)}%")


def producer(doc, c):
    doc.add_heading("Producer Agreement (Draft for Review)", level=1)
    doc.add_paragraph(f"Artist: {c.get('artist', 'TBD')}")
    doc.add_paragraph(f"Producer: {c.get('producer', 'TBD')}")
    numbered(doc, f"Fee / advance: {c.get('fee', '$TBD')} (recoupable against royalty points)")
    numbered(doc, f"Royalty points: {c.get('points', '3-5')}% on artist's royalty from first recoupment")
    numbered(doc, "Master ownership: 100% to artist unless stated otherwise")
    numbered(doc, f"Composition split (if producer co-writes): {c.get('publishing_split', '50/50')}")
    numbered(doc, "Credit: 'Produced by [Producer]' on all formats")
    numbered(doc, "Recoupment & accounting: quarterly statements, audit rights")


def management(doc, c):
    doc.add_heading("Management Agreement (Draft for Review)", level=1)
    doc.add_paragraph(f"Artist: {c.get('artist', 'TBD')}")
    doc.add_paragraph(f"Manager: {c.get('manager', 'TBD')}")
    numbered(doc, f"Commission: {c.get('commission', '15-20')}% on {c.get('commission_base', 'gross (or modified gross)')}")
    numbered(doc, f"Term: initial {c.get('term', '12')} months with measurable objectives")
    numbered(doc, "Exclusions: funds not comissionable (e.g., recording advances designated for costs) — list explicitly")
    numbered(doc, "Sunset clause after termination:")
    for s in c.get("sunset", []):
        doc.add_paragraph(f"- {s}")
    numbered(doc, "Manager obligations: strategy, coordination of booking/publicist/lawyer/business manager")


def feature(doc, c):
    doc.add_heading("Feature / Collaboration Agreement (Draft for Review)", level=1)
    doc.add_paragraph(f"Lead artist: {c.get('lead', 'TBD')}")
    doc.add_paragraph(f"Featured artist: {c.get('featured', 'TBD')}")
    numbered(doc, f"Feature fee: {c.get('fee', '$TBD')} or royalty share: {c.get('royalty_share', 'TBD')}%")
    numbered(doc, f"Publishing split: {c.get('publishing_split', '50/50')}")
    numbered(doc, "Clearances: sample and third-party rights responsibility")
    numbered(doc, "Credit & release approval rights")


def build(c: dict) -> None:
    doc = Document()
    kind = c.get("contract_type", "split_sheet")
    doc.add_heading("DRAFT — for lawyer review only. Not legal advice.", level=2)
    {
        "split_sheet": split_sheet,
        "producer": producer,
        "management": management,
        "feature": feature,
    }.get(kind, split_sheet)(doc, c)

    doc.add_paragraph("")
    doc.add_paragraph("Terms and blanks marked TBD/0 must be completed and reviewed by qualified counsel before signing.")

    out_dir = "/workspace/out"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"{kind}_draft.docx")
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    build(load_input())
