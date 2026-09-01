#!/usr/bin/env python3
"""Generate a .docx press kit (EPK) from a JSON bundle.

Convention:
  - Read JSON from /workspace/in/epk.json (or first argv).
  - Write output to /workspace/out/EPK_<slug>.docx.
Run only via the execute_python sandbox. python-docx is preinstalled.
"""
import json
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def load_input() -> dict:
    path = sys.argv[1] if len(sys.argv) > 1 else "/workspace/in/epk.json"
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def add_kv(doc, label, value):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(str(value))


def build(epk: dict) -> None:
    doc = Document()
    title = epk.get("artist_name", "Artist")
    slug = epk.get("slug", title.lower().replace(" ", "-").replace("/", "-"))
    doc.add_heading(title, 0)

    tagline = epk.get("tagline")
    if tagline:
        p = doc.add_paragraph(tagline)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Long bio
    if epk.get("long_bio"):
        add_heading(doc, "Biography", 1)
        doc.add_paragraph(epk["long_bio"])

    # Short bio / one-sheet bullets
    if epk.get("short_bio"):
        add_heading(doc, "Short Bio", 1)
        doc.add_paragraph(epk["short_bio"])

    if epk.get("one_sheet"):
        add_heading(doc, "One Sheet", 1)
        for item in epk["one_sheet"]:
            doc.add_paragraph(item, style="List Bullet")

    # Stats
    if epk.get("stats"):
        add_heading(doc, "Key Stats", 1)
        for stat in epk["stats"]:
            if isinstance(stat, dict):
                add_kv(doc, stat.get("label", "Stat"), stat.get("value", ""))
            else:
                doc.add_paragraph(str(stat))

    # Key tracks
    if epk.get("tracks"):
        add_heading(doc, "Key Tracks", 1)
        for t in epk["tracks"]:
            if isinstance(t, dict):
                line = f"{t.get('title', '')} — {t.get('note', '')}".strip(" —")
            else:
                line = str(t)
            doc.add_paragraph(line, style="List Bullet")

    # Quotes
    if epk.get("quotes"):
        add_heading(doc, "Quotes", 1)
        for q in epk["quotes"]:
            doc.add_paragraph(f'"{q}"', style="List Bullet")

    # Socials + contacts
    add_heading(doc, "Connect", 1)
    for social in epk.get("socials", []):
        add_kv(doc, social.get("network", "Social"), social.get("handle", ""))
    for contact in epk.get("contacts", []):
        add_kv(doc, contact.get("role", "Contact"), contact.get("email", ""))

    out_dir = "/workspace/out"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"EPK_{slug}.docx")
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    build(load_input())
