#!/usr/bin/env python3
"""Generate a cue sheet (.docx + .xlsx) for a production.

A cue sheet lists every musical cue in a film/TV/game production so PROs can
pay performance royalties. Input JSON: /workspace/in/cue_sheet.json with:
  production, episodes or segments: [{title, duration, source}],
  cues: [{cue, title, composer, publisher, usage (BG/VF/instrumental),
          duration_sec, isrc}]
Output: /workspace/out/cue_sheet.docx + cue_sheet.xlsx

Run only via the execute_python sandbox.
"""
import json
import os
import sys

import pandas as pd
from docx import Document
from docx.shared import Pt


def load_input() -> dict:
    path = sys.argv[1] if len(sys.argv) > 1 else "/workspace/in/cue_sheet.json"
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def build_docx(data: dict, cues: list) -> str:
    doc = Document()
    doc.add_heading(f"Cue Sheet — {data.get('production', 'Production')}", 0)
    for seg in data.get("segments", []):
        doc.add_paragraph(f"{seg.get('title', '')} ({seg.get('duration', '')}) — source: {seg.get('source', '')}")
    doc.add_heading("Cues", 1)
    table = doc.add_table(rows=1, cols=7)
    hdr = table.rows[0].cells
    for i, h in enumerate(["Cue", "Title", "Composer", "Publisher", "Usage", "Duration (s)", "ISRC"]):
        hdr[i].text = h
    for c in cues:
        row = table.add_row().cells
        for i, val in enumerate([
            c.get("cue", ""), c.get("title", ""), c.get("composer", ""),
            c.get("publisher", ""), c.get("usage", ""), str(c.get("duration_sec", "")),
            c.get("isrc", ""),
        ]):
            row[i].text = val

    out_dir = "/workspace/out"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "cue_sheet.docx")
    doc.save(out_path)
    return out_path


def main() -> None:
    data = load_input()
    cues = data.get("cues", [])
    df = pd.DataFrame(cues)

    docx_path = build_docx(data, cues)
    out_dir = "/workspace/out"
    xlsx_path = os.path.join(out_dir, "cue_sheet.xlsx")
    with pd.ExcelWriter(xlsx_path, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="cues", index=False)

    total = df["duration_sec"].sum() if not df.empty else 0
    print(f"Saved {docx_path}")
    print(f"Saved {xlsx_path} — {len(cues)} cues, {total:.0f}s total music")


if __name__ == "__main__":
    main()
